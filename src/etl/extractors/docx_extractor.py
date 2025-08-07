"""
DOCX document extractor
"""

import glob
from typing import Iterator, Dict, Any
import pandas as pd
from pathlib import Path
from docx import Document
from .base import BaseExtractor


class DOCXExtractor(BaseExtractor):
    """Extractor for DOCX documents"""
    
    def validate_config(self) -> bool:
        """Validate DOCX extractor configuration"""
        required_keys = ['paths']
        for key in required_keys:
            if key not in self.config:
                raise ValueError(f"Missing required config key: {key}")
        return True
    
    def extract(self) -> Iterator[pd.DataFrame]:
        """
        Extract text content from DOCX files
        
        Yields:
            DataFrame with extracted text content
        """
        self.validate_config()
        
        paths = self.config['paths']
        
        for path_pattern in paths:
            files = glob.glob(path_pattern)
            
            if not files:
                self.logger.warning(f"No DOCX files found matching pattern: {path_pattern}")
                continue
            
            for file_path in files:
                self.logger.info(f"Extracting text from DOCX: {file_path}")
                
                try:
                    text_content = self._extract_docx_text(file_path)
                    
                    if text_content:
                        # Create DataFrame with extracted content
                        df = pd.DataFrame({
                            'content': [text_content],
                            'file_name': [Path(file_path).name],
                            'file_path': [file_path],
                            '_source_type': ['docx']
                        })
                        yield df
                    else:
                        self.logger.warning(f"No text extracted from: {file_path}")
                        
                except Exception as e:
                    self.logger.error(f"Error extracting from DOCX {file_path}: {e}")
                    continue
    
    def _extract_docx_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = []
            if paragraphs:
                all_text.extend(paragraphs)
            if tables_text:
                all_text.append("\n--- Tables ---")
                all_text.extend(tables_text)
            
            return "\n".join(all_text)
            
        except Exception as e:
            self.logger.error(f"Error reading DOCX file {file_path}: {e}")
            raise